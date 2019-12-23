import docx
import os
from docx.shared import Inches
from PIL import Image
from PIL import ImageFont
from PIL import ImageDraw
import sys

ANCHOX24 = 906
ALTURAX24 = 453

ANCHOX12 = 453
ALTURAX12 = 453


def root_folder():
    ret = ''
    path = sys.argv[0].split('/')

    for k in range(0, len(path) - 1):
        ret = ret + path[k] + '\\'

    return ret

def res_img(img, tipo):
    if(tipo == 'X12'):
        formateada = img.resize((ANCHOX12, ALTURAX12))
    else:
        formateada = img.resize((ANCHOX24, ALTURAX24))
    return formateada

def abrir(name):
    a = Image.open("images\\" + name)
    return a

def getRGB(color):

     if color == 'black':
          ret = (0,0,0)
     if color == 'white':
          ret = (255,255,255)
     if color == 'red':
          ret = (255,0,0)
     if color == 'green':
          ret = (0,255,0)
     if color == 'blue':
          ret = (0,0,255)
     if color == 'yellow':
          ret = (255,255,0)
     
     return ret
       
def Draw(imagen,tipo,lang,color):
     
     draw = ImageDraw.Draw(imagen)
     fontMayus = ImageFont.truetype("calibri.ttf", 22)
     font = ImageFont.truetype("calibri.ttf", 16)
     colour = getRGB (color)

     if (tipo == 'X24'):

        if (lang == 'SPA'):
                draw.text((486, 428), "CASTELLANO",colour, font=fontMayus)

        draw.text((802, 433), "es copia", colour, font=font)

     else:
          if (lang == 'SPA'):
               draw.text((30, 428), "CASTELLANO",colour, font=fontMayus)

          draw.text((377, 434), "es copia", colour, font=font)

          imagen.save('temporal.jpg')


          images = [Image.open(x) for x in ['temporal.jpg', 'temporal.jpg']]
          widths, heights = zip(*(i.size for i in images))

          total_width = sum(widths)
          max_height = max(heights)

          new_im = Image.new('RGB', (total_width, max_height))

          x_offset = 0

          for im in images:
               new_im.paste(im, (x_offset,0))
               x_offset += im.size[0]

          imagen = new_im
          

     imagen.save('temporal.jpg')

def load_image (doc):
          doc.add_picture('temporal.jpg', width=Inches(9.44882), height=Inches(4.72441))
          doc.add_picture('temporal.jpg', width=Inches(9.44882), height=Inches(4.72441))
          doc.add_picture('temporal.jpg', width=Inches(9.44882), height=Inches(4.72441))
          return doc

def save_to_docx(name, tipo, lang, color):
     img = abrir(name)
     img = res_img(img, tipo)
     Draw(img, tipo, lang, color)
     doc = docx.Document('resources\\PlantillaA3.docx')
     doc = load_image (doc)
     doc.save('Covers\\' + '%s.docx' % name.split('.')[0])


os.chdir(r"D:\MIS_COSAS\MyProjects\Phyton\Dvd_cover")

for file in os.listdir(root_folder() + "images\\" ):
    print("Insert the parameters for " + file + "\n (Type) (lang) (Color):\t")
    parameters = input().split()
    save_to_docx(file, parameters[0] , parameters [1] , parameters [2])

os.remove("temporal.jpg")
exit()